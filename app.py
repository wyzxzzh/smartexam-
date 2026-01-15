import streamlit as st
import openai
import pypandoc
import os
import tempfile
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

st.set_page_config(
    page_title="SmartExam - 智能出题系统",
    page_icon="📝",
    layout="wide"
)

def set_font(run, font_name_cn, font_name_en, size, bold=False):
    run.font.name = font_name_en
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name_cn)

def set_paragraph_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=24):
    paragraph.alignment = alignment
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing = Pt(line_spacing)
    paragraph_format.space_after = Pt(6)
    paragraph_format.space_before = Pt(6)

def create_formatted_word(content, subject, difficulty):
    doc = Document()
    
    section = doc.sections[0]
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)
    section.left_margin = Pt(90)
    section.right_margin = Pt(90)
    
    title = doc.add_paragraph()
    set_paragraph_format(title, WD_ALIGN_PARAGRAPH.CENTER)
    title_run = title.add_run(f"{subject}练习题")
    set_font(title_run, "黑体", "SimHei", 22, True)
    
    subtitle = doc.add_paragraph()
    set_paragraph_format(subtitle, WD_ALIGN_PARAGRAPH.CENTER)
    subtitle_run = subtitle.add_run(f"难度：{difficulty}")
    set_font(subtitle_run, "楷体", "KaiTi", 14)
    
    doc.add_paragraph()
    
    lines = content.split('\n')
    current_section_title = None
    
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        
        if line.startswith('#'):
            level = line.count('#')
            text = line.lstrip('#').strip()
            
            if level == 2:
                p = doc.add_paragraph()
                set_paragraph_format(p, WD_ALIGN_PARAGRAPH.LEFT)
                run = p.add_run(text)
                set_font(run, "黑体", "SimHei", 14, True)
                current_section_title = text
            elif level == 3:
                p = doc.add_paragraph()
                set_paragraph_format(p, WD_ALIGN_PARAGRAPH.LEFT)
                run = p.add_run(text)
                set_font(run, "黑体", "SimHei", 12, True)
            continue
        
        if line.startswith(('A.', 'B.', 'C.', 'D.', 'A、', 'B、', 'C、', 'D、')):
            p = doc.add_paragraph()
            set_paragraph_format(p)
            option_letter = line[0]
            option_text = line[2:] if line[1] in ['.', '、'] else line[1:]
            run = p.add_run(f"{option_letter}. ")
            set_font(run, "宋体", "Times New Roman", 12, True)
            run = p.add_run(option_text)
            set_font(run, "宋体", "Times New Roman", 12)
        elif line[0].isdigit() and line[1] in ['.', '、']:
            p = doc.add_paragraph()
            set_paragraph_format(p)
            run = p.add_run(line)
            set_font(run, "宋体", "Times New Roman", 12)
        else:
            p = doc.add_paragraph()
            set_paragraph_format(p)
            run = p.add_run(line)
            set_font(run, "宋体", "Times New Roman", 12)
    
    return doc

def apply_word_formatting(docx_path, subject, difficulty):
    doc = Document(docx_path)
    
    for paragraph in doc.paragraphs:
        if not paragraph.text.strip():
            continue
        
        if paragraph.style.name.startswith('Heading'):
            level = int(paragraph.style.name[-1])
            for run in paragraph.runs:
                if level == 1:
                    set_font(run, "黑体", "SimHei", 22, True)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif level == 2:
                    set_font(run, "黑体", "SimHei", 14, True)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif level == 3:
                    set_font(run, "黑体", "SimHei", 12, True)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            text = paragraph.text.strip()
            
            if text and text[0] in ['A', 'B', 'C', 'D'] and len(text) > 1 and text[1] in ['.', '、']:
                option_letter = text[0]
                option_content = text[2:] if text[1] in ['.', '、'] else text[1:]
                
                paragraph.clear()
                run1 = paragraph.add_run(f"{option_letter}. ")
                set_font(run1, "宋体", "Times New Roman", 12, True)
                run2 = paragraph.add_run(option_content)
                set_font(run2, "宋体", "Times New Roman", 12)
            else:
                for run in paragraph.runs:
                    set_font(run, "宋体", "Times New Roman", 12)
        
        set_paragraph_format(paragraph)
    
    section = doc.sections[0]
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)
    section.left_margin = Pt(90)
    section.right_margin = Pt(90)
    
    copyright_para = doc.add_paragraph()
    set_paragraph_format(copyright_para, WD_ALIGN_PARAGRAPH.CENTER)
    copyright_run = copyright_para.add_run("© 海盐县钟战华")
    set_font(copyright_run, "宋体", "Times New Roman", 10)
    
    return doc

with st.sidebar:
    st.title("⚙️ 参数设置")
    
    st.divider()
    
    subject = st.selectbox(
        "学科",
        ["语文", "数学", "英语", "科学", "历史与社会"],
        index=1
    )
    
    difficulty = st.selectbox(
        "难度",
        ["基础 (C)", "提升 (B)", "培优 (A)"],
        index=1
    )
    
    st.divider()
    
    st.subheader("题量设置")
    
    single_choice_count = st.number_input(
        "单选题数量",
        min_value=0,
        max_value=20,
        value=5,
        step=1
    )
    
    fill_blank_count = st.number_input(
        "填空题数量",
        min_value=0,
        max_value=20,
        value=3,
        step=1
    )
    
    short_answer_count = st.number_input(
        "简答题数量",
        min_value=0,
        max_value=10,
        value=1,
        step=1
    )
    
    st.divider()
    
    creativity = st.slider(
        "创意度",
        min_value=0.0,
        max_value=1.0,
        value=0.5,
        step=0.1,
        help="0.0 为保守模式，1.0 为创意模式"
    )
    
    st.divider()

st.title("📚 SmartExam - 智能出题系统")
st.markdown("基于课本内容的初中练习题自动生成工具")

st.divider()

input_text = st.text_area(
    "请输入课文内容或知识点",
    height=300,
    placeholder="在此粘贴教材文本或知识点，系统将根据内容自动生成练习题..."
)

if st.button("🚀 生成练习题", type="primary", use_container_width=True):
    if not input_text.strip():
        st.error("请输入课文内容或知识点")
        st.stop()
    
    if single_choice_count == 0 and fill_blank_count == 0 and short_answer_count == 0:
        st.error("请至少设置一种题型的数量")
        st.stop()
    
    with st.spinner("正在生成练习题，请稍候..."):
        try:
            client = openai.OpenAI(
                api_key=st.secrets["deepseek_api_key"],
                base_url="https://api.deepseek.com"
            )
            
            prompt = f"""你是一位资深的初中教师。请根据以下教材内容或知识点，生成一套标准化的练习题。

学科：{subject}
难度：{difficulty}

题量要求：
- 单选题：{single_choice_count} 题
- 填空题：{fill_blank_count} 题
- 简答题：{short_answer_count} 题

创意度：{creativity}（0.0 为保守模式，1.0 为创意模式）

教材内容/知识点：
{input_text}

输出格式要求（严格遵守）：

1. 整体结构：
   - 第一行：## 一、选择题
   - 第二行：## 二、填空题
   - 第三行：## 三、简答题
   - 第四行：## 参考答案

2. 题目编号格式：
   - 使用"1."、"2."、"3."的格式
   - 每道题之间空一行

3. 选项格式：
   - 使用"A."、"B."、"C."、"D."的格式
   - 选项字母后加空格，然后是选项内容
   - 每个选项独占一行

4. 数学公式格式：
   - 所有数学符号、公式必须使用 LaTeX 格式
   - 必须包裹在单美元符号 $ 中（例如 $x^2$）
   - 不要使用 \[ \] 块级公式，全部使用行内公式

5. 题目内容要求：
   - 题目简洁明了，符合初中生认知水平
   - 避免使用过于复杂的表述
   - 确保题目与教材内容紧密相关

6. 参考答案格式：
   - 使用"1. xxx"的格式
   - 答案准确简洁
   - 每题答案独占一行
   - 每题答案后必须提供详细的解析
   - 解析格式：在答案后另起一行，使用"解析："开头
   - 解析要详细说明解题思路和步骤

示例格式：
## 一、选择题
1. 已知关于 $x$ 的一元二次方程 $x^2 - 2kx + k^2 - 1 = 0$ 有两个不相等的实数根，则实数 $k$ 的取值范围是（ ）
A. $k > -1$
B. $k \ge 0$
C. $k > 0$
D. $k > 1$

2. ...

## 二、填空题
1. 若 $a > 0$，则 $a$ 的相反数是______。

2. ...

## 三、简答题
1. 请简述一元二次方程的求根公式。

2. ...

## 参考答案
1. D
解析：一元二次方程有两个不相等的实数根，判别式 $\Delta > 0$，即 $(-2k)^2 - 4 \times 1 \times (k^2 - 1) > 0$，化简得 $4k^2 - 4k^2 + 4 > 0$，即 $4 > 0$，恒成立。但题目要求有两个不相等的实数根，所以 $k^2 - 1 \ne 0$，即 $k \ne \pm 1$。又因为 $k^2 - 1 = 0$ 时方程有一个实数根，所以 $k^2 - 1 > 0$，即 $k > 1$ 或 $k < -1$。结合选项，选 D。

2. ...

请严格按照以上格式生成练习题："""

            response = client.chat.completions.create(
                model="deepseek-chat",
                messages=[
                    {"role": "system", "content": "你是一位资深的初中教师，擅长根据教材内容出题。"},
                    {"role": "user", "content": prompt}
                ],
                temperature=creativity,
                max_tokens=4000
            )
            
            generated_content = response.choices[0].message.content
            
            st.success("✅ 练习题生成成功！")
            
            st.markdown("### 📄 生成的练习题")
            st.markdown(generated_content)
            
            with tempfile.TemporaryDirectory() as temp_dir:
                md_file = os.path.join(temp_dir, "temp_output.md")
                docx_file = os.path.join(temp_dir, "output.docx")
                
                with open(md_file, "w", encoding="utf-8") as f:
                    f.write(generated_content)
                
                pypandoc.convert_file(
                    md_file,
                    "docx",
                    outputfile=docx_file
                )
                
                doc = apply_word_formatting(docx_file, subject, difficulty)
                
                final_docx = os.path.join(temp_dir, "final_output.docx")
                doc.save(final_docx)
                
                with open(final_docx, "rb") as f:
                    docx_bytes = f.read()
                
                st.download_button(
                    label="📥 下载 Word 文档",
                    data=docx_bytes,
                    file_name=f"练习题_{subject}_{difficulty}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            
            st.download_button(
                label="📥 下载 Markdown 文档",
                data=generated_content,
                file_name=f"练习题_{subject}_{difficulty}.md",
                mime="text/markdown",
                use_container_width=True
            )
            
        except openai.AuthenticationError:
            st.error("API Key 验证失败，请检查你的 API Key 是否正确")
        except openai.APIError as e:
            st.error(f"API 调用失败：{str(e)}")
        except Exception as e:
            st.error(f"生成过程中出现错误：{str(e)}")

st.divider()
st.caption("© 海盐县钟战华")
st.caption("Powered by DeepSeek V3 & Pandoc")
